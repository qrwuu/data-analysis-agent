# -*- coding: utf-8 -*-
"""
file_parser.py — Parse docx/xlsx files into structured knowledge records.

Two paths:
  Structured   (xlsx with recognised column headers) → direct field mapping
  Unstructured (docx, or xlsx without template headers) → LLM extraction

The parser is stateless: pass in an OpenAI-compatible client + model name,
exactly the same objects BusinessAgent receives from get_llm_client().
"""
from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Any

log = logging.getLogger(__name__)

import pandas as pd

# ── Template detection ────────────────────────────────────────────────────────

# Minimum number of recognised header keywords to classify a sheet as structured
_TEMPLATE_THRESHOLD = 2

_METRIC_HEADERS: dict[str, str] = {
    # canonical field  : accepted column aliases (lowercase)
    "name":         {"name", "名称", "指标名", "指标", "metric"},
    "alias":        {"alias", "别名", "别称", "aliases"},
    "definition":   {"definition", "定义", "desc", "description", "说明"},
    "sql_template": {"sql", "sql_template", "sql模板", "query"},
    "notes":        {"notes", "备注", "note", "remark", "说明2"},
}

_RULE_HEADERS: dict[str, str] = {
    "rule_id":     {"rule_id", "规则id", "id", "rule"},
    "description": {"description", "描述", "说明", "desc"},
    "condition":   {"condition", "条件", "断言", "assert"},
    "severity":    {"severity", "严重程度", "level", "等级"},
}

_NOTE_HEADERS: dict[str, str] = {
    "topic":   {"topic", "主题", "话题", "subject"},
    "content": {"content", "内容", "text", "正文"},
    "tags":    {"tags", "标签", "tag", "关键词"},
}

_ALL_TEMPLATE_SETS = [
    ("metrics",        _METRIC_HEADERS),
    ("business_rules", _RULE_HEADERS),
    ("context_notes",  _NOTE_HEADERS),
]


def _detect_template(columns: list[str]) -> tuple[str, int]:
    """Return (table_type, match_count) for the best-matching template."""
    norm = {c.strip().lower() for c in columns}
    best_type, best_count = "", 0
    for table_type, header_map in _ALL_TEMPLATE_SETS:
        count = sum(
            1 for aliases in header_map.values()
            if norm & aliases
        )
        if count > best_count:
            best_type, best_count = table_type, count
    return best_type, best_count


def _map_columns(df: pd.DataFrame, header_map: dict[str, set]) -> list[str]:
    """Return a column-name mapping: {canonical_field: actual_df_column}."""
    norm_to_actual = {c.strip().lower(): c for c in df.columns}
    mapping: dict[str, str | None] = {}
    for field, aliases in header_map.items():
        hit = next((norm_to_actual[a] for a in aliases if a in norm_to_actual), None)
        mapping[field] = hit
    return mapping


def _df_to_structured(df: pd.DataFrame, table_type: str) -> list[dict]:
    """Map a structured DataFrame to a list of preview records."""
    header_map = dict(_ALL_TEMPLATE_SETS)[table_type]
    col_map = _map_columns(df, header_map)
    records = []
    for _, row in df.iterrows():
        rec = {"table": table_type}
        for field, col in col_map.items():
            rec[field] = str(row[col]).strip() if col and pd.notna(row.get(col)) else ""
        # Minimal validation: skip completely empty rows
        values = [v for k, v in rec.items() if k != "table"]
        if any(v for v in values):
            records.append(rec)
    return records


# ── Text extraction ───────────────────────────────────────────────────────────

def _extract_docx_text(filepath: str) -> str:
    from docx import Document
    doc = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # Also extract table cells
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _df_to_text(df: pd.DataFrame) -> str:
    """Convert an unstructured DataFrame to readable text for LLM."""
    lines = []
    for _, row in df.iterrows():
        cells = [str(v).strip() for v in row if pd.notna(v) and str(v).strip()]
        if cells:
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_text(filepath: str) -> str:
    """Extract readable text from a supported knowledge source file.

    Used by the RAG indexer after the user confirms import.  It deliberately
    keeps sheet names and table-like rows so retrieved chunks retain source
    context even when the original file was semi-structured.
    """
    path = Path(filepath)
    ext = path.suffix.lower()
    if ext == ".docx":
        return _extract_docx_text(filepath)
    if ext in (".xlsx", ".xls"):
        xl = pd.ExcelFile(filepath)
        parts = []
        for sheet in xl.sheet_names:
            df = xl.parse(sheet)
            if df.empty:
                continue
            parts.append(f"[Sheet: {sheet}]\n{_df_to_text(df)}")
        return "\n\n".join(parts)
    raise ValueError(f"Unsupported file type: {ext}")


# ── LLM extraction ────────────────────────────────────────────────────────────

_EXTRACT_PROMPT = """你是业务知识整理助手。请从以下文本中提取所有业务指标定义、业务规则和背景知识。

严格按照以下 JSON 格式输出，不要输出任何其他内容：
{
  "metrics": [
    {
      "table": "metrics",
      "name": "指标英文或中文名称（必填）",
      "alias": "别名，逗号分隔，没有则留空",
      "definition": "指标的业务定义",
      "sql_template": "计算该指标的 SQL 模板，没有则留空",
      "notes": "口径说明、注意事项等"
    }
  ],
  "business_rules": [
    {
      "table": "business_rules",
      "rule_id": "规则的英文标识符（如 retention_sanity）",
      "description": "规则的中文描述",
      "condition": "违反该规则的条件描述（自然语言）",
      "severity": "error 或 warning"
    }
  ],
  "context_notes": [
    {
      "table": "context_notes",
      "topic": "主题",
      "content": "背景知识内容",
      "tags": "相关标签，逗号分隔"
    }
  ]
}

规则：
- 无法提取的字段留空字符串，不要编造
- 同一指标只保留一条记录
- 如果文本中没有某类知识，对应数组返回 []

---文本开始---
{text}
---文本结束---"""


_CHUNK_SIZE = 12000  # chars per chunk


def _parse_raw_to_records(raw: str) -> list[dict]:
    """Parse LLM JSON response into a flat list of records."""
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.MULTILINE)
    raw = re.sub(r"\s*```\s*$", "", raw, flags=re.MULTILINE)
    raw = raw.strip()

    data = None
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        log.debug("[file_parser] json.loads failed: %s", e)

    if data is None:
        match = re.search(r"\{[\s\S]*\}", raw)
        if match:
            try:
                data = json.loads(match.group())
            except json.JSONDecodeError as e:
                log.debug("[file_parser] fallback json.loads failed: %s", e)

    if data is None or not isinstance(data, dict):
        log.warning("[file_parser] could not extract JSON. raw=%r", raw[:300])
        return []

    records = []
    for key in ("metrics", "business_rules", "context_notes"):
        items = data.get(key, [])
        if not isinstance(items, list):
            continue
        for item in items:
            if not isinstance(item, dict):
                continue
            item["table"] = key
            records.append(item)
    return records


def _llm_call_once(chunk: str, client, model: str) -> list[dict]:
    """Call LLM for a single text chunk using streaming to avoid read timeouts."""
    prompt = _EXTRACT_PROMPT.replace("{text}", chunk)
    log.info("[file_parser] LLM call model=%s chunk_len=%d", model, len(chunk))
    try:
        stream = client.with_options(max_retries=0, timeout=120).chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=131072,
            stream=True,
        )
        parts = []
        for chunk_ev in stream:
            delta = chunk_ev.choices[0].delta if chunk_ev.choices else None
            if delta and delta.content:
                parts.append(delta.content)
        raw = "".join(parts).strip()
    except Exception as e:
        log.error("[file_parser] LLM API call failed: %s", e)
        raise RuntimeError(f"LLM 调用失败: {e}") from e
    log.info("[file_parser] LLM response (first 300 chars): %r", raw[:300])
    return _parse_raw_to_records(raw)


def _llm_extract(text: str, client, model: str) -> list[dict]:
    """Split text into chunks and extract records from each, then merge."""
    # Split on blank lines to avoid cutting mid-sentence
    paragraphs = re.split(r"\n{2,}", text.strip())

    chunks: list[str] = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) + 2 > _CHUNK_SIZE and current:
            chunks.append(current.strip())
            current = para
        else:
            current = (current + "\n\n" + para).strip() if current else para
    if current:
        chunks.append(current.strip())

    if not chunks:
        return []

    log.debug("[file_parser] split into %d chunk(s) for model=%s", len(chunks), model)

    all_records: list[dict] = []
    seen_names: set[str] = set()   # deduplicate by metric name / rule_id / topic

    for i, chunk in enumerate(chunks):
        log.debug("[file_parser] processing chunk %d/%d", i + 1, len(chunks))
        recs = _llm_call_once(chunk, client, model)
        for rec in recs:
            # Dedup key varies by table
            table = rec.get("table", "")
            if table == "metrics":
                key = rec.get("name", "").strip().lower()
            elif table == "business_rules":
                key = rec.get("rule_id", "").strip().lower()
            else:
                key = rec.get("topic", "").strip().lower()
            uid = f"{table}::{key}"
            if key and uid in seen_names:
                continue
            if key:
                seen_names.add(uid)
            all_records.append(rec)

    log.debug("[file_parser] total records extracted: %d", len(all_records))
    return all_records


# ── Public API ────────────────────────────────────────────────────────────────

def parse_file(filepath: str, client, model: str) -> dict[str, Any]:
    """Parse a docx or xlsx file into a preview list of knowledge records.

    Returns:
      {
        "format":  "structured" | "unstructured",
        "preview": [ {table, field1, field2, ...}, ... ]
      }

    The caller (api/knowledge.py) returns this directly to the frontend.
    After the user reviews and edits, POST /api/knowledge/confirm with the
    final list → KnowledgeBase.bulk_insert().
    """
    path = Path(filepath)
    ext = path.suffix.lower()

    if ext in (".xlsx", ".xls"):
        # Try every sheet; use the first one that looks structured
        xl = pd.ExcelFile(filepath)
        structured_records: list[dict] = []
        unstructured_texts: list[str] = []

        for sheet in xl.sheet_names:
            df = xl.parse(sheet)
            if df.empty:
                continue
            table_type, confidence = _detect_template(list(df.columns))
            if confidence >= _TEMPLATE_THRESHOLD:
                structured_records.extend(_df_to_structured(df, table_type))
            else:
                unstructured_texts.append(f"[Sheet: {sheet}]\n{_df_to_text(df)}")

        # If we found any structured sheets, return them directly.
        # Unstructured sheets in the same file are also sent to LLM.
        preview: list[dict] = list(structured_records)
        if unstructured_texts:
            combined = "\n\n".join(unstructured_texts)
            preview.extend(_llm_extract(combined, client, model))

        fmt = "structured" if structured_records and not unstructured_texts else "unstructured"
        if structured_records and unstructured_texts:
            fmt = "mixed"

        return {"format": fmt, "preview": preview}

    elif ext == ".docx":
        text = _extract_docx_text(filepath)
        preview = _llm_extract(text, client, model)
        return {"format": "unstructured", "preview": preview}

    else:
        raise ValueError(f"Unsupported file type: {ext}. Please upload .xlsx or .docx")
