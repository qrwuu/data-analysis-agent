import logging
log = logging.getLogger(__name__)
import os
import datetime
import zipfile
import tempfile
from typing import Any, Dict, List, Optional, Tuple


def export_to_report(
    title: str,
    sections: List[Dict[str, Any]],
    filepath: str,
    chart_htmls: Optional[List[str]] = None,
) -> Tuple[str, str]:
    """
    Generate a Word (.docx) report and optionally bundle charts as a ZIP.

    Parameters
    ----------
    title       : Report title shown at the top of the document.
    sections    : Ordered list of {heading: str, content: str | list[dict]}.
                  str content → paragraph; list[dict] content → table.
    filepath    : Absolute path for the .docx file (directory must exist).
    chart_htmls : Optional list of Plotly HTML strings from the session.
                  When provided, returns a .zip containing the .docx + chart HTML files.

    Returns
    -------
    (result_path, download_name)
      result_path   : absolute path of the file to serve (zip or docx)
      download_name : filename the browser will see
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise ImportError("python-docx 未安装，请运行: pip install python-docx") from exc

    doc = Document()

    # ── Title block ────────────────────────────────────────────────────
    doc.add_heading(title, level=0)
    ts_para = doc.add_paragraph(
        f"生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    ts_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    ts_para.runs[0].font.size = Pt(10)
    doc.add_paragraph()  # spacer

    # ── Sections ───────────────────────────────────────────────────────
    for section in sections:
        heading = section.get("heading", "")
        content = section.get("content", "")

        if heading:
            doc.add_heading(heading, level=1)

        if isinstance(content, str) and content.strip():
            # Split on double newlines to preserve paragraph breaks
            for para in content.split("\n\n"):
                para = para.strip()
                if para:
                    doc.add_paragraph(para)

        elif isinstance(content, list) and content:
            cols = list(content[0].keys())
            tbl = doc.add_table(rows=1, cols=len(cols))
            tbl.style = "Table Grid"
            hdr_cells = tbl.rows[0].cells
            for i, col in enumerate(cols):
                hdr_cells[i].text = str(col)
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.bold = True
            for row_data in content:
                row_cells = tbl.add_row().cells
                for i, col in enumerate(cols):
                    row_cells[i].text = str(row_data.get(col, ""))
            doc.add_paragraph()  # spacer after table

    # ── Chart reference section ─────────────────────────────────────────
    if chart_htmls:
        doc.add_heading("图表附件", level=1)
        doc.add_paragraph(
            f"本报告共包含 {len(chart_htmls)} 张交互式图表，"
            "已随报告一同打包为 HTML 文件，请用浏览器打开查看。"
        )
        for i in range(len(chart_htmls)):
            doc.add_paragraph(f"• chart_{i + 1:02d}.html")

    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    doc.save(filepath)

    # ── Bundle into ZIP when charts are present ─────────────────────────
    if chart_htmls:
        zip_path = filepath.replace(".docx", ".zip")
        zip_name = os.path.basename(zip_path)
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.write(filepath, os.path.basename(filepath))
            for i, html in enumerate(chart_htmls):
                zf.writestr(f"chart_{i + 1:02d}.html", html)
        return zip_path, zip_name

    doc_name = os.path.basename(filepath)
    return filepath, doc_name
